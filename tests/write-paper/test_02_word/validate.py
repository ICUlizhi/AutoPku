#!/usr/bin/env python3
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).parent.parent.parent))
from conftest import check_file_exists

paper_docx = Path("test00/财务报表分析/论文/paper.docx")
results = []

# 1. paper.docx 存在
ok, msg = check_file_exists(paper_docx)
results.append(msg)

if not paper_docx.exists():
    # 若不存在（AI 尚未执行），测试框架应标记为需要手动执行
    results.append("INFO: paper.docx 不存在，需先执行 run.sh 中的 AI Agent 步骤")
    failed = sum(1 for r in results if r.startswith("FAIL"))
    for r in results:
        print(r)
    sys.exit(0 if failed == 0 else 1)

# 2. 使用 python-docx 验证文档内容
try:
    from docx import Document
    from docx.oxml.ns import qn
except ImportError:
    results.append("FAIL: python-docx 未安装，无法验证 Word 文档")
    failed = sum(1 for r in results if r.startswith("FAIL"))
    for r in results:
        print(r)
    sys.exit(1)

doc = Document(str(paper_docx))

# 3. core_properties 作者为学生姓名（非 python-docx）
props = doc.core_properties
author = props.author or ""
if author.lower() == "python-docx":
    results.append(f"FAIL: core_properties.author 为 '{author}'，未替换为学生姓名")
elif author.strip() == "":
    results.append("FAIL: core_properties.author 为空，未设置学生姓名")
else:
    results.append(f"OK: core_properties.author = '{author}'（已替换）")

# 4. 标题存在且不为空
title_found = False
for para in doc.paragraphs[:10]:
    text = para.text.strip()
    if text and len(text) > 5:
        title_found = True
        results.append(f"OK: 标题/首段内容存在: '{text[:40]}...'")
        break
if not title_found:
    results.append("FAIL: 文档标题/正文为空")

# 5. 中文字体检查
# 检查至少有一个 run 使用了中文字体（黑体/宋体/楷体）
zh_font_found = False
zh_fonts = {"黑体", "宋体", "楷体", "SimHei", "SimSun", "KaiTi", "FangSong"}
for para in doc.paragraphs:
    for run in para.runs:
        rfonts = run._element.rPr.rFonts
        if rfonts is not None:
            east_asia = rfonts.get(qn("w:eastAsia"))
            if east_asia and any(zf in str(east_asia) for zf in zh_fonts):
                zh_font_found = True
                results.append(f"OK: 中文字体已设置: {east_asia}")
                break
    if zh_font_found:
        break
if not zh_font_found:
    results.append("FAIL: 未找到中文字体设置（eastAsia 字体）")

# 6. 摘要存在
abstract_found = False
for para in doc.paragraphs:
    if "摘要" in para.text:
        abstract_found = True
        results.append("OK: 摘要段落存在")
        break
if not abstract_found:
    results.append("FAIL: 未找到摘要段落")

# 7. 参考文献存在
ref_found = False
for para in doc.paragraphs:
    if "参考文献" in para.text or "[1]" in para.text:
        ref_found = True
        results.append("OK: 参考文献段落存在")
        break
if not ref_found:
    results.append("FAIL: 未找到参考文献段落")

failed = sum(1 for r in results if r.startswith("FAIL"))
for r in results:
    print(r)

sys.exit(0 if failed == 0 else 1)
